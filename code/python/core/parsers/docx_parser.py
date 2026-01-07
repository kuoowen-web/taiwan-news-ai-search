# Copyright (c) 2025 Microsoft Corporation.
# Licensed under the MIT License

"""
DOCX file parser using python-docx.
"""

from typing import Dict, Any
from .base_parser import BaseParser
from misc.logger.logging_config_helper import get_configured_logger

logger = get_configured_logger("docx_parser")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. DOCX parsing will not be available.")


class DOCXParser(BaseParser):
    """Parser for DOCX files using python-docx."""

    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a DOCX file and extract text content.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Dictionary containing extracted text and metadata
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Install it with: pip install python-docx")

        try:
            doc = Document(file_path)

            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            table_texts.append(cell.text)

            # Combine all text
            full_text = '\n\n'.join(paragraphs)
            if table_texts:
                full_text += '\n\n' + '\n'.join(table_texts)

            metadata = {
                'num_paragraphs': len(paragraphs),
                'num_tables': len(doc.tables),
                'total_chars': len(full_text)
            }

            logger.info(f"Parsed DOCX: {len(paragraphs)} paragraphs, {len(full_text)} characters")

            return {
                'text': full_text,
                'metadata': metadata
            }

        except Exception as e:
            logger.exception(f"Failed to parse DOCX: {str(e)}")
            raise

    def supports_file_type(self, file_extension: str) -> bool:
        """Check if this parser supports DOCX files."""
        return file_extension.lower() == '.docx'
